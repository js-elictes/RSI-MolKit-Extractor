#!/usr/bin/env python3

""" A simple program that extracts geometries and thermodynamic data of optimized Gaussian jobs. 
Put all your results into one folder and generate an Excel table with all thermodynamics or an .xyz file for supplementary materials. 
Place in the directory of your files and run. """

import os
import logging
import re
from openpyxl import Workbook
import docx
from docx.shared import Cm, Pt

# Constants and logging
__version__ = "1.7 -- 10.04.2025"
Hartree_to_kJ = 2625.4996394799
error_rate = 0
logging.basicConfig(level=logging.INFO)


def do_not_overwrite(path):
    n, e = os.path.splitext(path)
    i = 1
    while os.path.exists(path):
        path = f"{n} ({i}){e}"
        i += 1
    return path


def get_input(prompt, options):
    while True:
        choice = input(prompt).lower()
        if choice in options:
            return options[choice]
        logging.error(" -Select a valid option-")


def input_prompt():
    # Prompts the user for output preferences (Excel, Docs, or XYZ).
    print(f"\n  -- SuperJoel {__version__} by Jonáš Schröder --\n")
    option = get_input("Output an Excel, Docs, or XYZ file? [Excel/Docs/XYZ] : ",
                       {"excel": "excel", "e": "excel",
                        "docs": "docs", "d": "docs",
                        "xyz": "xyz", "x": "xyz"})
    print("\n  -- processing -- \n")
    return option


def export_relevant(log_file, option):
    global error_rate
    try:
        frq_header, ngeom = None, ""
        with open(log_file, 'r') as imported_file:
            content = imported_file.read()
        for i in re.finditer(r'---*\n (#.*?)---*', content, re.DOTALL):
            if "freq" in "".join(j.strip() for j in i.group(0)).lower():
                frq_header = i.group(1).replace("\n ", "")
                frq_header_pos = i.span()
        end_frq_pos = (re.search(r'Normal termination', content[frq_header_pos[1]:]).span()[1] +
                       frq_header_pos[1])
        frq_calc = content[frq_header_pos[0]:end_frq_pos]
        thermochem = " " + re.search(r'(Zero-point correction= .*?\n) \n', frq_calc, re.DOTALL).group(1)
        geom = re.findall(r' *Standard orientation: *\n -*\n.*?-*\n -*\n(.*?) -{10}', frq_calc, re.DOTALL)[-1]
        for i in geom.splitlines():
            num, atom, atype, x, y, z = i.split()
            ngeom = ngeom + " ".join([atom, " ", x, " ", y, " ", z]) + "\n"
        logging.info(f" Processing file  :  {log_file}")
        if option == "variables":
            charge = int(re.search(r'-?\d+', re.search(r'Charge = .*?(?= Multiplicity)', frq_calc).group(0)).group())
            mult = int(re.search(r'-?\d+', re.search(r'Multiplicity = .*?\n', frq_calc).group(0)).group())
            thermochem_vals = [float(x) for x in re.findall(r'-?\d*\.\d+|-?\d+', thermochem)]
            thermochem_vals = [val for i, val in enumerate(thermochem_vals) if i not in (1, 2, 3, 5)]
            imag_values = [float(x) for x in re.findall(r'-?\d*\.\d+|-?\d+', 
                              re.search(r'Low frequencies ---.*?\n', frq_calc).group(0))]
            imag = "0" if all(abs(val) < 30 for val in imag_values) else imag_values[0]
            E_tot = thermochem_vals[1] - thermochem_vals[0]
            E_ok, H_298k, G_298k = thermochem_vals[1:]
            return frq_header, charge, mult, imag, E_tot, E_ok, H_298k, G_298k, ngeom
        else:
            chrgandmult = re.search(r'Charge = .*? Multiplicity = .*?\n', frq_calc).group(0)
            lowfrqs = "".join(re.findall(r'Low frequencies ---.*?\n', frq_calc))
            outstr = "\n\n".join([log_file, frq_header, chrgandmult, thermochem, lowfrqs])
            return outstr
    except Exception as e:
        logging.error(f" Critical failure :  {log_file}")
        error_rate += 1
        outstr = "\n\n".join([log_file, "⚠️ This file encountered an Error\n"])
        if option == "variables":
            return None
        else:
            return outstr


def create_excel_output(log_files):
    wb = Workbook()
    ws = wb.active
    header = [
        "File name", "Header", "Charge", "Multiplicity", "Imag",
        "E-tot (Hartree)", "E-tot / rel (kJ/mol)", "E-ok (Hartree)",
        "E-ok / rel (kJ/mol)", "H-298k (Hartree)", "H-298k / rel (kJ/mol)",
        "G-298k (Hartree)", "G-298k / rel (kJ/mol)"
    ]
    ws.append(header)
    dataset = [export_relevant(f, "variables") for f in log_files]
    Energs = [data[4] if data else float('inf') for data in dataset]
    smallest = Energs.index(min(Energs))
    for idx, data in enumerate(dataset):
        log_file = log_files[idx].replace(".log", "")
        if not data:
            ws.append([log_file, '⚠️ This file encountered an Error'])
            continue
        frqheader, charge, mult, imag, E_tot, E_ok, H_298k, G_298k, geom = data
        Etotrel, Eokrel, H298rel, G298rel = [
            round(abs(dataset[smallest][i] - data[i]) * Hartree_to_kJ, 1)
            for i in range(4, 8)
        ]
        row = [
            log_file, frqheader, charge, mult, imag, E_tot, Etotrel,
            E_ok, Eokrel, H_298k, H298rel, G_298k, G298rel
        ]
        ws.append(row)
    return wb


def create_word_output(log_files):
    doc = docx.Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    style = doc.styles['Normal']
    style.paragraph_format.space_before = style.paragraph_format.space_after = Cm(0)
    style.font.name, style.font.size = 'Arial', Pt(11)
    for log_file in log_files:
        lines = export_relevant(log_file, "string").split("\n")
        para = doc.add_paragraph(lines[0])
        run = para.runs[0]
        run.bold, run.font.size = True, Pt(14)
        for line in lines[1:]:
            doc.add_paragraph(line)
        doc.add_paragraph("")
    return doc


def create_xyz_output(log_files):
    merged_geometries = []
    for log_file in log_files:
        data = export_relevant(log_file, "variables")
        if not data:
            logging.error(f" Skipping {log_file}: export_relevant did not extract data.")
            continue
        else:
            frq_header, charge, mult, imag, E_tot, E_ok, H_298k, G_298k, ngeom = data
        coord_lines = [line for line in ngeom.splitlines() if line.strip() != ""]
        comment = f"{log_file} | E(HF)={E_tot:.6f} | E(0K)={E_ok:.6f} | Imag={imag} | Charge={charge} | Multiplicity={mult}"
        merged_geometries.append({'source': log_file, 'atoms': coord_lines, 'comment': comment})
    if not merged_geometries:
        return None
    merged_string = ""
    for geom in merged_geometries:
        count = len(geom['atoms'])
        merged_string += f"{count}\n"
        merged_string += f"{geom['comment']}\n"
        for line in geom['atoms']:
            merged_string += f"{line}\n"
    return merged_string


if __name__ == "__main__":
    option = input_prompt()
    log_files = [f for f in os.listdir() if f.endswith('.log')]
    if option == "excel":
        export_file = do_not_overwrite("SuperJoel_Excel_Output.xlsx")
        create_excel_output(log_files).save(export_file)
        print(f"\n Excel file created: {os.path.abspath(export_file)}")
    elif option == "docs":
        export_file = do_not_overwrite("SuperJoel_Word_Output.docx")
        create_word_output(log_files).save(export_file)
        print(f"\n Word file created: {os.path.abspath(export_file)}")
    elif option == "xyz":
        export_file = do_not_overwrite("SuperJoel_XYZ_Output.xyz")
        with open(export_file, "w") as f:
            f.write(create_xyz_output(log_files))
        print(f"\n XYZ file created: {os.path.abspath(export_file)}")
        
    print(f"\n  -- Finished -- {error_rate} out of {len(log_files)} files encountered an Error --\n")
    print(r"""             __..--''``---....___   _..._    __
   /// //_.-'    .-/";  `        ``<._  ``.''_ `. / // /
  ///_.-' _..--.'_    \                    `( ) ) // //
  / (_..-' // (< _     ;_..__               ; `' / ///
   / // // //  `-._,_)' // / ``--...____..-' /// / //\
""")
    quit()
